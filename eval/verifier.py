"""Deterministic verifier tests (pytest). Run: pytest eval/verifier.py -v

These check the parts with exactly one correct answer: conservation, sign
constraints, Monte Carlo reproducibility, and that the agent's reported numbers
match the simulator (no fabrication).
"""
import copy

import numpy as np
import pytest

from agent import skills
from agent.orchestrator import WaitCostAgent, TierViolation
from model.simulate import simulate
from model.states import STATES

PARAMS_PATH = "config/params.yaml"


@pytest.fixture
def params():
    return skills.fetch_hud_data(PARAMS_PATH)


def test_population_conserved_without_inflow(params):
    p = copy.deepcopy(params)
    p["inflow"] = {}                       # turn off external inflow
    sc = skills.make_scenario(p, "test", budget=0.0)
    df = simulate(p, sc)
    totals = df[STATES].sum(axis=1)
    assert np.allclose(totals, totals.iloc[0], rtol=1e-6), "population must be conserved"


def test_no_negative_stocks(params):
    sc = skills.make_scenario(params, "test", delay=0, budget=20.0)
    df = simulate(params, sc)
    assert (df[STATES].to_numpy() >= -1e-6).all(), "stocks must stay non-negative"


def test_cumulative_cost_nonnegative_and_monotonic(params):
    sc = skills.make_scenario(params, "test", budget=10.0)
    df = simulate(params, sc)
    cc = df["cum_cost"].to_numpy()
    assert (cc >= -1e-6).all(), "cumulative cost must be non-negative"
    assert np.all(np.diff(cc) >= -1e-6), "cumulative cost must be non-decreasing"


def test_montecarlo_reproducible(params):
    sc = skills.make_scenario(params, "test", budget=10.0)
    a = skills.run_simulation(params, sc, n_mc=50)["mc_final"]
    b = skills.run_simulation(params, sc, n_mc=50)["mc_final"]
    assert np.array_equal(a, b), "fixed seed must give identical Monte Carlo output"


def test_delay_costs_more_than_acting_now(params):
    """Face validity: with a beneficial intervention, delaying should not be cheaper."""
    now = float(simulate(params, skills.make_scenario(params, "now", delay=0, budget=15.0))["cum_cost"].iloc[-1])
    delay = float(simulate(params, skills.make_scenario(params, "delay", delay=5, budget=15.0))["cum_cost"].iloc[-1])
    assert delay >= now - 1e-6, "delaying a beneficial intervention should cost >= acting now"


def test_agent_reports_match_simulator(params):
    """The agent must not invent numbers: its reported P50 equals the simulator's."""
    agent = WaitCostAgent(PARAMS_PATH, memory_path="MEMORY.md", max_auto_tier=1)
    result = agent.answer("What if we wait 5 years?", out_dir="outputs")
    # recompute act-now P50 independently
    sc = skills.make_scenario(params, "Act now", delay=0,
                              budget=result["plan"]["budget_musd"])
    expected = skills.run_simulation(params, sc)["final_cum_cost_p50"]
    reported = result["runs"]["act_now"]["final_cum_cost_p50"]
    assert reported == pytest.approx(expected, rel=1e-9), "reported number must match simulator"


def test_tier2_requires_approval(params):
    """Action-Tier guard: a Tier-2 skill must raise without human approval."""
    agent = WaitCostAgent(PARAMS_PATH, max_auto_tier=1)
    with pytest.raises(TierViolation):
        agent._check_tier("optimize_allocation", approve=False)


# --- Learned inflow model (the real AI component) --------------------------
from model import inflow_model as im   # noqa: E402

PANEL = "data/coc_panel.csv"


def test_inflow_model_reproducible():
    """The learned pipeline is deterministic: same data -> same metric/prediction."""
    a = im.train_and_calibrate(PANEL, target_coc="CA-600")
    b = im.train_and_calibrate(PANEL, target_coc="CA-600")
    assert a["loo_r2"] == pytest.approx(b["loo_r2"])
    assert a["predicted_rate_per_1k"] == pytest.approx(b["predicted_rate_per_1k"])


def test_inflow_model_has_heldout_metric():
    """Model must report a real cross-validated (leave-one-CoC-out) metric."""
    rep = im.train_and_calibrate(PANEL, target_coc="CA-600")
    assert -1.0 <= rep["loo_r2"] <= 1.0
    assert rep["loo_mae"] > 0
    # held-out R^2 must be strictly below in-sample R^2 (i.e., honestly evaluated)
    assert rep["loo_r2"] < rep["insample_r2"] + 1e-9


def test_shap_values_are_exact_additive():
    """Exact-SHAP property: per-feature attributions sum to prediction - E[f]."""
    rep = im.train_and_calibrate(PANEL, target_coc="CA-600")
    shap_sum = sum(d["shap"] for d in rep["shap_target"])
    gap = rep["predicted_rate_per_1k"] - rep["expected_value_rate_per_1k"]
    assert shap_sum == pytest.approx(gap, abs=1e-6)


def test_inflow_central_is_spm_anchored(params):
    """Central inflow must equal HUD SPM Measure 5 (first-time homeless / 12)."""
    spm_m5_monthly = 29818 / 12.0
    assert params["inflow"]["at_risk"] == pytest.approx(spm_m5_monthly, rel=0.01)
    assert params.get("inflow_uncertainty", {}).get("cv", 0) > 0


def test_ml_inflow_corroborates_spm():
    """Independent check: the ACS->PIT ML inflow agrees with HUD SPM M5 within 20%."""
    rep = im.train_and_calibrate(PANEL, target_coc="CA-600",
                                 spm_first_time_annual=29818, spm_pit=71320)
    cv = rep["spm_crossval"]
    assert cv is not None
    assert cv["agreement_pct_diff"] < 20.0, "ML and SPM inflow should agree within 20%"


def test_backtest_brackets_observed_2024(params):
    """Face validity: seeded on real 2023 PIT, the 12-month run brackets observed 2024."""
    from model.backtest import backtest
    r = backtest(params, n_mc=120)
    assert r["within_band"], "observed 2024 PIT must fall within the backtest P10-P90"
    assert r["abs_pct_error_p50"] < 15.0, "central backtest error should be modest"


# --- Bypass / data-sufficiency guard (responsible AI, enforced not just documented) ---
def test_bypass_passes_for_real_coc(params):
    """The real CA-600 calibration has enough support: the guard must NOT trip."""
    assert skills.check_data_support(params) is True


def test_bypass_declines_thin_data(params):
    """Thin homeless count -> the agent must decline rather than show bands."""
    p = copy.deepcopy(params)
    for s in ["sheltered", "unsheltered", "chronic_unsheltered"]:
        p["initial_population"][s] = 10
    with pytest.raises(skills.DataSufficiencyError):
        skills.check_data_support(p)


def test_bypass_declines_sub_coc(params):
    """Sub-CoC geography -> decline (source data can't support it)."""
    p = copy.deepcopy(params)
    p["meta"] = dict(p["meta"]); p["meta"]["sub_coc"] = True
    with pytest.raises(skills.DataSufficiencyError):
        skills.check_data_support(p)


# --- Multi-question intent layer -------------------------------------------
from agent import planner   # noqa: E402


def test_intent_classification():
    cases = {
        "What if we wait 3 years on a $15M program?": "cost_of_waiting",
        "How long can we afford to wait?": "break_even",
        "How much do we save by acting now?": "savings_now",
        "How many people will be homeless if we do nothing?": "outcome_at_horizon",
        "Is $15M or $50M better?": "compare_budgets",
        "Should we fund prevention or supportive housing?": "compare_mix",
        "Which assumption are we least sure about?": "sensitivity",
        "Who specifically will become homeless on 5th street?": "out_of_scope",
        # greetings + meta/capability questions route to the friendly path
        "Hello,": "greeting",
        "hi": "greeting",
        "thanks": "greeting",
        "good morning": "greeting",
        "what can you do?": "greeting",
        "help": "greeting",
        "who are you?": "greeting",
        # new analysis intents
        "What is the ROI on a $15M program?": "roi",
        "Is it worth the investment?": "roi",
        "How many people would acting now help?": "cost_per_person",
        "What is the cost per person helped?": "cost_per_person",
        "Which cities pay the most for waiting?": "regional",
        "Rank the cities by cost of inaction": "regional",
        "How confident are you in that number?": "uncertainty",
        "Can we trust this estimate?": "uncertainty",
        # in-scope but unmapped -> guide the user, don't guess
        "What is the capital of France?": "clarify",
    }
    for q, want in cases.items():
        assert planner.classify_intent(q) == want, f"{q!r} -> {planner.classify_intent(q)} (want {want})"


def test_compare_budgets_skill(params):
    out = skills.compare_budgets(params, [15.0, 50.0], delay=0, n_mc=20)
    assert len(out["budgets"]) == 2
    assert out["best_budget_musd"] in (15.0, 50.0)
    assert all(r["cum_cost_p50"] > 0 for r in out["budgets"])


def test_compare_mix_skill(params):
    out = skills.compare_mix(params, 50.0, delay=0, n_mc=20)
    assert len(out["mixes"]) == 4
    assert out["best_mix"] in {r["mix"] for r in out["mixes"]}


def test_agent_routes_intent_and_gives_direct_answer():
    agent = WaitCostAgent(PARAMS_PATH, memory_path="MEMORY.md", max_auto_tier=1)
    res = agent.answer("How much do we save by acting now?", out_dir="outputs")
    assert res["intent"] == "savings_now"
    assert res["direct_answer"] and "save" in res["direct_answer"].lower()


def test_agent_declines_out_of_scope():
    agent = WaitCostAgent(PARAMS_PATH, memory_path="MEMORY.md", max_auto_tier=1)
    res = agent.answer("Which family on 5th street will become homeless?", out_dir="outputs")
    assert res.get("declined") and res.get("out_of_scope")
    assert "runs" not in res


def test_agent_greets_without_running_engine():
    """A greeting/meta message returns a warm guidance card (greeting=True), is NOT
    framed as an out-of-scope refusal, and never runs the simulation engine."""
    agent = WaitCostAgent(PARAMS_PATH, memory_path="MEMORY.md", max_auto_tier=1)
    for q in ("Hello,", "what can you do?"):
        res = agent.answer(q, out_dir="outputs")
        assert res.get("declined") and res.get("greeting")
        assert not res.get("out_of_scope")
        assert "runs" not in res                    # engine never ran
        assert "cost of waiting" in res["reason"].lower()


def test_agent_clarifies_unmapped_in_scope():
    """An in-scope-but-unmapped question is guided (clarify), not silently answered
    as cost_of_waiting, and never runs the engine."""
    agent = WaitCostAgent(PARAMS_PATH, memory_path="MEMORY.md", max_auto_tier=1)
    res = agent.answer("What is the capital of France?", out_dir="outputs")
    assert res.get("declined") and res.get("greeting") and res.get("intent") == "clarify"
    assert "runs" not in res
    assert "not sure" in res["reason"].lower()


def test_agent_answers_new_intents():
    """ROI / cost-per-person / uncertainty produce sensible, number-bearing answers."""
    agent = WaitCostAgent(PARAMS_PATH, memory_path="MEMORY.md", max_auto_tier=1)
    roi = agent.answer("What is the ROI on a $15M program?", out_dir="outputs")
    assert roi["intent"] == "roi" and ("$" in roi["direct_answer"])
    cpp = agent.answer("How many people would acting now help?", out_dir="outputs")
    assert cpp["intent"] == "cost_per_person" and "people" in cpp["direct_answer"].lower()
    unc = agent.answer("How confident are you in that number?", out_dir="outputs")
    assert unc["intent"] == "uncertainty" and "%" in unc["direct_answer"]


def test_agent_regional_ranks_cities():
    """The regional intent ranks the cost of waiting across multiple cities."""
    from agent import skills
    reg = skills.regional_cost_of_waiting(50.0, 5, n_mc=40)
    assert len(reg["cities"]) >= 2
    vals = [c["cost_of_waiting_musd"] for c in reg["cities"]]
    assert vals == sorted(vals, reverse=True)        # sorted high -> low
    agent = WaitCostAgent(PARAMS_PATH, memory_path="MEMORY.md", max_auto_tier=1)
    res = agent.answer("Which cities pay the most for waiting?", out_dir="outputs")
    assert res["intent"] == "regional" and res["recommended_chart"] == "regional_waiting"


def test_viz_recommends_chart_matching_the_answer():
    """The recommended chart must present the quantity the text answer states."""
    from analysis.viz import VizAgent
    va = VizAgent()
    assert va.recommend("cost_of_waiting") == "cost_of_waiting"   # the waterfall, not the trajectory
    assert va.recommend("roi") == "roi"
    assert va.recommend("cost_per_person") == "people_helped"
    assert va.recommend("regional") == "regional_waiting"
    assert va.recommend("uncertainty") == "sensitivity_tornado"


def test_chart_numbers_match_agent_answer():
    """The cost-of-waiting waterfall, built with the same seeded Monte Carlo as the
    agent, must show the SAME headline figure as the text (the core viz-match fix)."""
    from api import payloads as P
    agent = WaitCostAgent(PARAMS_PATH, memory_path="MEMORY.md", max_auto_tier=1)
    res = agent.answer("What if we wait 3 years on a $15M program?", out_dir="outputs")
    agent_cow = round(res["comparison"]["cost_of_waiting"]["extra_cost_median"] / 1e6, 1)
    spec = P.chart_payload("cost_of_waiting", coc="CA-600", budget=15, delay=3)
    # The true waterfall's first bar is now the act-now total; the cost-of-waiting
    # slice is the "relative" step, kept addressable via the series' extra_cost field.
    assert spec["series"][0]["extra_cost"] == pytest.approx(agent_cow), "chart slice must equal the text"
    assert spec["series"][0]["measure"] == ["absolute", "relative", "total"]


def test_brief_cites_card_verdict_verbatim(monkeypatch):
    """Single source of truth: the Decision brief must QUOTE the Recommendation
    card's verdict sentence (decision['headline']) verbatim — never recompute or
    rephrase it — and frame the repeat as supporting evidence, so the two surfaces
    can never contradict each other. Forced to rule mode to be hermetic; the
    orchestrator's citation guarantee holds in the Gemma path too."""
    monkeypatch.setenv("WAITCOST_PLANNER", "rule")
    agent = WaitCostAgent(PARAMS_PATH, memory_path="MEMORY.md", max_auto_tier=1)
    res = agent.answer("What if we wait 3 years on a $15M program?", out_dir="outputs")
    verdict = res["decision"]["headline"]
    brief = res["brief_markdown"]
    assert verdict, "decision must expose the verdict sentence as headline"
    assert verdict in brief, "brief must quote the card's verdict sentence character-for-character"
    assert "This brief supports the recommendation above:" in brief, \
        "brief must frame the repeat as supporting evidence, not a second verdict"


# --- Feature ①: provenance on every number --------------------------------
def test_provenance_payload_covers_every_metric_family():
    """GET /provenance must return an entry for each metric family with a
    non-empty source + vintage (sourced from SOURCES_MANIFEST.md, not invented)."""
    from api import payloads as P
    prov = P.provenance_payload()
    for fam in ("homeless_counts", "economic_features", "flow_rates", "costs",
                "cost_of_waiting", "scenario", "equity", "model"):
        assert fam in prov, f"missing provenance family: {fam}"
        assert prov[fam]["source"].strip(), f"{fam} has empty source"
        assert prov[fam]["vintage"].strip(), f"{fam} has empty vintage"
        assert prov[fam]["label"].strip(), f"{fam} has empty label"
    # The real-data families must trace to the manifest, not be hard-coded here.
    assert "Census" in prov["economic_features"]["source"]
    assert "range" in prov["cost_of_waiting"]["note"].lower()


def test_gemma4_is_the_default_model():
    assert planner.GEMMA_MODEL == "gemma4:e2b"
    assert planner.resolve_model(["gemma4:e2b", "x"]) == "gemma4:e2b"


# --- Multi-CoC: the SAME model reused for other cities ---------------------
def test_build_params_reuses_model_for_other_city():
    from model.coc_registry import available_cocs, build_params_for_coc
    assert len(available_cocs()) >= 15
    p = build_params_for_coc("WA-500")          # Seattle/King County
    ip = p["initial_population"]
    # initial homeless compartments come from real PIT (sheltered + unshel + chronic = PIT total)
    assert ip["sheltered"] + ip["unsheltered"] + ip["chronic_unsheltered"] == 16868
    assert p["inflow"]["at_risk"] > 0           # inflow predicted by the model
    # simulate runs cleanly for the other city
    df = simulate(p, skills.make_scenario(p, "now", budget=15.0))
    assert (df[STATES].to_numpy() >= -1e-6).all()


def test_same_model_different_inflow_per_city():
    from model.coc_registry import build_params_for_coc
    # Same trained model, different real ACS inputs -> different predicted inflows.
    wa = build_params_for_coc("WA-500")["inflow"]["at_risk"]
    fl = build_params_for_coc("FL-600")["inflow"]["at_risk"]
    assert wa != fl


def test_agent_runs_for_other_city():
    from model.coc_registry import build_params_for_coc
    p = build_params_for_coc("IL-510")          # Chicago
    agent = WaitCostAgent(PARAMS_PATH, memory_path="MEMORY.md", max_auto_tier=1, params=p)
    res = agent.answer("What if we wait 3 years on a $15M program?", out_dir="outputs")
    assert "Chicago" in res["plan"]["question"] or res["intent"] == "cost_of_waiting"
    assert res["direct_answer"]


# --- Tool registry (function-calling catalog) + US-context retrieval -------
def test_tool_registry():
    from agent import tools
    s = tools.registry_summary()
    assert s["agents"] >= 1
    assert s["capabilities"] >= 9
    names = tools.capability_names()
    assert "retrieve_us_context" in names and "recommend_allocation" in names
    assert tools.tier_of("recommend_allocation") == 2   # human-gated


def test_retrieve_us_context_real():
    ctx = skills.retrieve_us_context("CA-600")
    assert ctx["indicators"]["homeless_pit_total"] == 71201
    assert ctx["indicators"]["median_home_value_usd"] > 0
    wa = skills.retrieve_us_context("WA-500")
    assert wa["indicators"]["homeless_pit_total"] == 16868
    import pytest as _pytest
    with _pytest.raises(ValueError):
        skills.retrieve_us_context("ZZ-999")


def test_city_context_intent_and_answer():
    # city_context now owns the NUMERIC snapshot phrasings; the narrative "tell me
    # about / overview / situation" phrasings route to the CityBriefAgent (city_situation).
    assert planner.classify_intent("Show me this city's housing cost-of-living profile") == "city_context"
    assert planner.classify_intent("What are the housing indicators?") == "city_context"
    agent = WaitCostAgent(PARAMS_PATH, memory_path="MEMORY.md", max_auto_tier=1)
    res = agent.answer("Show me this city's housing profile", out_dir="outputs")
    assert res["intent"] == "city_context"
    assert "homeless" in res["direct_answer"].lower() and "Los Angeles" in res["direct_answer"]


# --- Visualization agent (charts as function-calling specs) ----------------
def test_registry_has_four_agents():
    from agent import tools
    s = tools.registry_summary()
    assert s["agents"] == 4
    assert s["agent_names"] == ["analyst", "visualization", "city_brief", "decision"]
    assert s["charts"] >= 13
    assert "visualize" in tools.capability_names()


def test_viz_recommender_maps_intents():
    from analysis.viz import VizAgent
    va = VizAgent()
    assert va.recommend("compare_mix") == "mix_comparison"
    assert va.recommend("break_even") == "break_even_curve"
    assert va.recommend("sensitivity") == "sensitivity_tornado"


def test_viz_builds_every_chart_as_valid_spec():
    import json
    from analysis.viz import VizAgent, CHART_CATALOG
    va = VizAgent()
    for c in CHART_CATALOG:
        spec = va.build(c["name"], coc="CA-600", budget=15.0, delay=3, n_mc=40)
        assert spec["kind"] == c["kind"] and spec["series"], c["name"]
        json.dumps(spec)   # must be JSON-serializable for the API/frontend


def test_agent_attaches_recommended_chart():
    agent = WaitCostAgent(PARAMS_PATH, memory_path="MEMORY.md", max_auto_tier=1)
    res = agent.answer("Should we fund prevention or supportive housing?", out_dir="outputs")
    assert res["recommended_chart"] == "mix_comparison"


# --- Compound questions (tool-calling: one question -> a LIST of calls) -----
_META = {"meta": {"default_budget_musd": 10.0}}


def test_compound_budget_answers_every_value_no_silent_drop():
    """A two-budget question answers BOTH budgets — never silently drops one."""
    agent = WaitCostAgent(PARAMS_PATH, memory_path="MEMORY.md", max_auto_tier=1)
    res = agent.answer("What does it cost to wait 3 years on a $1M and a $15M program?", out_dir="outputs")
    assert res["intent"] == "cost_of_waiting"
    assert [c["args"]["budget_musd"] for c in res["plan"]["calls"]] == [1.0, 15.0]
    assert "$1M" in res["direct_answer"] and "$15M" in res["direct_answer"]
    assert [r["budget_musd"] for r in res["sweep"]["rows"]] == [1.0, 15.0]
    assert res["recommended_chart"] == "cost_of_waiting_by_budget"


def test_compound_rows_match_single_budget_runs():
    """Each swept budget equals a separate single-budget /ask (determinism across
    the two code paths)."""
    agent = WaitCostAgent(PARAMS_PATH, memory_path="MEMORY.md", max_auto_tier=1)
    sweep = agent.answer("Cost of waiting 3 years on a $1M and a $15M program?", out_dir="outputs")["sweep"]
    by_b = {r["budget_musd"]: r["cost_of_waiting"]["extra_cost_median"] for r in sweep["rows"]}
    for b, q in [(1.0, "What if we wait 3 years on a $1M program?"),
                 (15.0, "What if we wait 3 years on a $15M program?")]:
        a2 = WaitCostAgent(PARAMS_PATH, memory_path="MEMORY.md", max_auto_tier=1)
        single = a2.answer(q, out_dir="outputs")["comparison"]["cost_of_waiting"]["extra_cost_median"]
        assert by_b[b] == pytest.approx(single, rel=1e-9)


def test_compare_budgets_not_reinterpreted_as_sweep():
    """'Is $15M or $50M better?' stays compare_budgets (one call)."""
    p = planner.plan("Is $15M or $50M the better budget?", _META)
    assert p["intent"] == "compare_budgets" and len(p["calls"]) == 1


def test_single_budget_path_unchanged():
    p = planner.plan("What if we wait 3 years on a $15M program?", _META)
    assert p["intent"] == "cost_of_waiting" and len(p["calls"]) == 1
    assert p["budget_musd"] == 15.0


def test_safety_rail_collapses_to_single_out_of_scope_call():
    p = planner.plan("Which family at $1M and $15M is evicted on 5th street?", _META)
    assert p["intent"] == "out_of_scope"
    assert len(p["calls"]) == 1 and p["calls"][0]["tool"] == "out_of_scope"


def test_duplicate_budgets_deduped_to_single_call():
    p = planner.plan("What does it cost to wait 3 years on a $15M and a $15M program?", _META)
    assert len(p["calls"]) == 1


def test_too_many_budgets_capped_with_note():
    p = planner.plan("Cost of waiting 3 years on $1M, $2M, $3M, $4M, $5M, $6M and $7M programs?", _META)
    assert len(p["calls"]) == 6   # planner.MAX_CALLS
    assert any("cap" in n.lower() or "first" in n.lower() for n in p["coverage_notes"])


# --- Retrieval tools (cited, offline, no engine) ---------------------------
def test_concept_qa_is_cited_and_skips_engine():
    agent = WaitCostAgent(PARAMS_PATH, memory_path="MEMORY.md", max_auto_tier=1)
    res = agent.answer("What is rapid re-housing?", out_dir="outputs")
    assert res["intent"] == "concept_qa"
    assert "runs" not in res                       # engine never ran
    assert res["sources"] and res["sources"][0]["url"]
    assert res["label"].startswith("General context")


def test_data_lookup_reports_real_vintage_no_engine():
    agent = WaitCostAgent(PARAMS_PATH, memory_path="MEMORY.md", max_auto_tier=1)
    res = agent.answer("What data is this based on?", out_dir="outputs")
    assert res["intent"] == "data_lookup"
    assert "runs" not in res
    assert "2024" in res["direct_answer"] and res["sources"]


# --- Equity / disparity lens (population-level, no individual profiling) ----
def test_equity_analysis_real_disproportionality():
    from analysis.equity import equity_analysis
    a = equity_analysis("CA-600")
    assert a["homeless_total"] == 71201
    top = a["most_overrepresented"]
    assert "Black" in top["group"] and top["factor"] > 3.0   # ~4x over-represented (real)
    # every group has a population share and an unsheltered rate
    for g in a["groups"]:
        assert g["population_share_pct"] >= 0 and 0 <= g["unsheltered_rate_pct"] <= 100


def test_equity_intent_not_out_of_scope():
    # population-level race questions are IN scope; individual ones are not
    assert planner.classify_intent("What are the racial disparities in homelessness?") == "equity"
    assert planner.classify_intent("Who is most affected by race?") == "equity"
    assert planner.classify_intent("Name the person on 5th street") == "out_of_scope"


def test_agent_answers_equity_for_la():
    agent = WaitCostAgent(PARAMS_PATH, memory_path="MEMORY.md", max_auto_tier=1)
    res = agent.answer("Show me the racial equity disparities", out_dir="outputs")
    assert res["intent"] == "equity"
    assert "over-represented" in res["direct_answer"] and res["recommended_chart"] == "equity_disparity"


def test_equity_charts_build():
    from analysis.viz import VizAgent
    va = VizAgent()
    for name in ("equity_disparity", "equity_unsheltered"):
        spec = va.build(name, coc="CA-600")
        assert spec["series"] and spec["kind"] == "bar"


def test_agent_declines_gracefully_on_thin_data(tmp_path):
    """End-to-end: the orchestrator returns a decline (no fabricated bands)."""
    import yaml
    p = skills.fetch_hud_data(PARAMS_PATH)
    p["meta"]["sub_coc"] = True
    cfg = tmp_path / "thin.yaml"
    cfg.write_text(yaml.safe_dump(p))
    agent = WaitCostAgent(str(cfg), memory_path=str(tmp_path / "MEM.md"), max_auto_tier=1)
    result = agent.answer("What if we wait 3 years?", out_dir=str(tmp_path))
    assert result.get("declined") is True
    assert "runs" not in result


def test_learned_inflow_widens_bands(params):
    """The learned inflow band must propagate: removing it must shrink the P10-P90."""
    sc = skills.make_scenario(params, "test", budget=15.0)
    with_band = skills.run_simulation(params, sc)
    p = copy.deepcopy(params)
    p.pop("inflow_uncertainty", None)            # turn off the learned inflow noise
    without = skills.run_simulation(p, sc)
    span_with = with_band["final_cum_cost_p90"] - with_band["final_cum_cost_p10"]
    span_without = without["final_cum_cost_p90"] - without["final_cum_cost_p10"]
    assert span_with > span_without, "learned inflow band must add uncertainty"


def test_data_vintage_is_not_synthetic(params):
    """Calibration must use real data, not the synthetic placeholder."""
    assert "synthetic" not in params["meta"]["data_vintage"].lower()


def test_brief_surfaces_learned_model_and_shap():
    """The AI reasoning must be visible in the brief: model, held-out metric, SHAP."""
    agent = WaitCostAgent(PARAMS_PATH, memory_path="MEMORY.md", max_auto_tier=1)
    result = agent.answer("What if we wait 3 years?", out_dir="outputs")
    md = result["brief_markdown"].lower()
    assert "learned inflow predictor" in md
    assert "shap" in md
    assert "leave-one-coc-out" in md or "r²" in md or "r^2" in md


# ===========================================================================
# FINALIZE: Gemma-default hardening + the four new features
# ===========================================================================

# --- Task A: Gemma-default planner with automatic, invisible fallback -------
def test_resolve_model_prefers_then_falls_back():
    """resolve_model: exact tag > same family > any Gemma > None."""
    assert planner.resolve_model([planner.GEMMA_MODEL, "x"]) == planner.GEMMA_MODEL
    # family fallback (asked e2b, only e4b pulled) -> same gemma family
    fam = planner.GEMMA_MODEL.split(":")[0]
    assert planner.resolve_model([f"{fam}:e4b"]).split(":")[0] == fam
    # any locally-installed Gemma keeps the demo live
    assert planner.resolve_model(["gemma4:e2b"]) == "gemma4:e2b"
    # nothing usable -> None (Ollama down / no Gemma)
    assert planner.resolve_model([]) is None
    assert planner.resolve_model(["llama3:8b"]) is None


def test_gemma_available_reflects_install(monkeypatch):
    monkeypatch.setattr(planner, "_installed_models", lambda: [])
    assert planner.gemma_available() is False
    monkeypatch.setattr(planner, "_installed_models", lambda: [planner.GEMMA_MODEL])
    assert planner.gemma_available() is True


def test_auto_planner_falls_back_when_ollama_fails(params, monkeypatch):
    """Default (auto) mode must degrade SILENTLY to rules when the LLM call dies —
    identical output shape, no exception."""
    monkeypatch.setenv("WAITCOST_PLANNER", "auto")
    monkeypatch.setattr(planner, "_installed_models", lambda: [planner.GEMMA_MODEL])

    def boom(*a, **k):
        raise RuntimeError("ollama unreachable")

    monkeypatch.setattr(planner, "_ollama_generate", boom)
    pl = planner.plan("What if we wait 5 years on a $15M program?", params)
    assert pl["planner"] == "rule_based_fallback"
    assert pl["intent"] == "cost_of_waiting" and pl["delay_years"] == 5


def test_auto_planner_skips_gemma_when_down(params, monkeypatch):
    """auto must not even attempt Gemma (no 20s stall) when Ollama is unreachable."""
    monkeypatch.setenv("WAITCOST_PLANNER", "auto")
    monkeypatch.setattr(planner, "_installed_models", lambda: [])

    def fail(*a, **k):
        raise AssertionError("must not call Ollama when no model is installed")

    monkeypatch.setattr(planner, "_ollama_generate", fail)
    assert planner.plan("How long can we wait?", params)["planner"] == "rule_based_fallback"


def test_forced_gemma_raises_when_down(params, monkeypatch):
    """`gemma` (forced) must surface the error rather than silently degrade."""
    monkeypatch.setenv("WAITCOST_PLANNER", "gemma")

    def boom(*a, **k):
        raise RuntimeError("ollama unreachable")

    monkeypatch.setattr(planner, "_ollama_generate", boom)
    with pytest.raises(Exception):
        planner.plan("What if we wait 5 years?", params)


# --- Feature ③: Gemma-written brief, number-guarded -------------------------
def test_number_guard_logic():
    facts = "- cost: $12.3M\n- savings: $4,500,000"
    assert planner.numbers_are_grounded("Waiting costs $12.3M over 10 years.", facts)
    assert planner.numbers_are_grounded("Saves $4,500,000.", facts)
    # a fabricated figure must be rejected
    assert not planner.numbers_are_grounded("It will cost $999.9M.", facts)


def test_narrate_brief_rejects_fabricated_numbers(monkeypatch):
    """Critical guard: if Gemma emits an unseen figure, narrate_brief returns None
    so the caller uses the deterministic brief."""
    monkeypatch.setenv("WAITCOST_PLANNER", "gemma")
    facts = {"cost_of_waiting_median": "$12.3M",
             "disclaimer": "informs not decides"}
    monkeypatch.setattr(planner, "_ollama_generate",
                        lambda *a, **k: "The cost of waiting is $999.9M over 10 years.")
    assert planner.narrate_brief(facts) is None
    # a grounded memo is accepted verbatim
    monkeypatch.setattr(planner, "_ollama_generate",
                        lambda *a, **k: "## Decision memo\nWaiting costs $12.3M. Informs, not decides.")
    out = planner.narrate_brief(facts)
    assert out and "12.3" in out


def test_agent_tags_brief_author_deterministic():
    """With the deterministic planner (test default), the brief author is deterministic."""
    agent = WaitCostAgent(PARAMS_PATH, memory_path="MEMORY.md", max_auto_tier=1)
    res = agent.answer("What if we wait 3 years?", out_dir="outputs")
    assert res["brief_author"] == "deterministic"
    assert res["planner"] == "rule_based_fallback"


# --- API: a TestClient over the FastAPI bridge ------------------------------
def _client():
    from fastapi.testclient import TestClient
    from api.main import app
    return TestClient(app)


def _parse_sse(text):
    events = []
    import json as _json
    for block in text.split("\n\n"):
        if not block.strip():
            continue
        ev = data = None
        for line in block.splitlines():
            if line.startswith("event:"):
                ev = line.split(":", 1)[1].strip()
            elif line.startswith("data:"):
                data = line.split(":", 1)[1].strip()
        events.append((ev, _json.loads(data) if data else None))
    return events


# --- Feature ①: streaming trajectory equals the synchronous answer ----------
def test_ask_stream_ends_with_result_equal_to_sync():
    c = _client()
    q = "What if we wait 3 years on a $15M program?"
    sync = c.post("/ask", json={"question": q}).json()
    r = c.post("/ask/stream", json={"question": q})
    assert r.status_code == 200
    events = _parse_sse(r.text)
    assert any(e == "step" for e, _ in events), "must stream at least one step"
    assert events[-1][0] == "result", "stream must end with a result event"
    streamed = events[-1][1]
    # determinism: the streamed numbers equal the synchronous /ask numbers
    assert (streamed["runs"]["act_now"]["final_cum_cost_p50"]
            == pytest.approx(sync["runs"]["act_now"]["final_cum_cost_p50"]))
    # steps carry a readable label + tier badge
    step = next(d for e, d in events if e == "step")
    assert step["label"] and "tier" in step and "status" in step


# --- Feature ②: PDF / Word export -------------------------------------------
def test_brief_export_docx_contains_headline():
    import io
    import re
    from docx import Document
    c = _client()
    q = "What if we wait 3 years on a $15M program?"
    sync = c.post("/ask", json={"question": q}).json()
    d = c.get("/brief/export", params={"format": "docx", "question": q})
    assert d.status_code == 200 and len(d.content) > 0
    assert "officedocument" in d.headers["content-type"]
    doc = Document(io.BytesIO(d.content))
    text = " ".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            text += " " + " ".join(cell.text for cell in row.cells)
    fig = re.search(r"\$\d[\d,]*\.?\d*M", sync["direct_answer"])
    assert fig, "sync answer should contain a $ figure"
    assert fig.group(0).replace(",", "") in text.replace(",", ""), "headline figure must appear in docx"


def test_brief_export_pdf_is_pdf():
    c = _client()
    p = c.get("/brief/export", params={"format": "pdf",
                                       "question": "What if we wait 3 years on a $15M program?"})
    assert p.status_code == 200 and len(p.content) > 0
    assert p.content[:4] == b"%PDF"


# --- Feature ④: two-city comparison -----------------------------------------
def test_compare_cities_returns_two_blocks_and_delta():
    c = _client()
    r = c.post("/compare-cities", json={"coc_a": "CA-600", "coc_b": "IL-510"})
    assert r.status_code == 200
    j = r.json()
    assert j["a"] and j["b"], "both city results present"
    assert j["a"]["intent"] == "cost_of_waiting" and j["b"]["intent"] == "cost_of_waiting"
    assert "cost_of_waiting_musd" in j["delta"]
    assert isinstance(j["delta"]["rate_per_1k"], (int, float))
